"""The response document — what actually goes to the procurement authority.

python-docx has no API for right-to-left text, so the two attributes that make
an Arabic document readable (``w:bidi`` on the paragraph, ``w:rtl`` on the run)
are set directly on the underlying XML. Without them Word lays Arabic out
left-to-right: the glyphs are correct, the line order is not, and punctuation
lands on the wrong end of every sentence. A submitted tender that looks like
that is a lost bid.
"""

from __future__ import annotations

import io
from typing import Final

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from app.exports.models import ExportBundle, ExportRow
from app.ingestion.normalizers.arabic import arabic_ratio

MEDIA_TYPE: Final = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

#: A font with real Arabic coverage. Word substitutes silently when a font
#: lacks glyphs, and the substitute is rarely one that shapes Arabic well.
_ARABIC_FONT: Final = "Arial"
_LATIN_FONT: Final = "Calibri"

_TITLE_AR: Final = "الرد على كراسة الشروط والمواصفات"
_TITLE_EN: Final = "Tender Response"
_PREPARED_AR: Final = "أُعد بواسطة"
_PREPARED_EN: Final = "Prepared by"


#: Share of Arabic letters above which a paragraph is laid out right-to-left.
#: Low on purpose: an Arabic sentence quoting "ISO/IEC 27001:2022" and a
#: certificate number is mostly Latin characters by count, and laying it out
#: left-to-right puts its full stop on the wrong end.
_RTL_THRESHOLD: Final = 0.15


def build_response_docx(bundle: ExportBundle) -> bytes:
    """Render the approved answers, in the tender's own order.

    Direction is decided **per paragraph**, from the text itself, rather than
    once for the whole document from the workspace language. A GCC tender is
    routinely Arabic prose with English standard references, and answers come
    back in either script inside one workspace — so a document-level flag gets
    it wrong for every paragraph in the other language, which is exactly the
    case a whole-document test never catches.
    """
    document = Document()
    _set_default_font(document)

    title = _TITLE_AR if bundle.is_rtl else _TITLE_EN
    _paragraph(document, title, bundle.is_rtl, style="Title")
    prepared = _PREPARED_AR if bundle.is_rtl else _PREPARED_EN
    _paragraph(
        document,
        f"{prepared}: {bundle.tenant_name} — {bundle.workspace_name}",
        bundle.is_rtl,
    )
    _paragraph(document, bundle.generated_at.strftime("%Y-%m-%d %H:%M UTC"), bundle.is_rtl)
    document.add_page_break()

    for row in bundle.answered_rows:
        _write_requirement(document, row)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ------------------------------------------------------------------ internals
def _write_requirement(document, row: ExportRow) -> None:
    heading = f"{row.requirement_ref} — {row.requirement_text}".strip(" —")
    # The requirement is quoted from the tender and the answer is written back
    # to it; they are frequently in different scripts, so each is measured
    # separately rather than inheriting one verdict.
    _paragraph(document, heading, is_rtl_text(heading), style="Heading 2")
    for block in (row.answer_text or "").split("\n"):
        if block.strip():
            _paragraph(document, block.strip(), is_rtl_text(block))


def _paragraph(document, text: str, rtl: bool, *, style: str | None = None) -> Paragraph:
    paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        _mark_paragraph_rtl(paragraph)
    run = paragraph.add_run(text)
    if rtl:
        _mark_run_rtl(run)
    return paragraph


def is_rtl_text(text: str) -> bool:
    """Whether this specific string should be laid out right-to-left."""
    return arabic_ratio(text) >= _RTL_THRESHOLD


def _mark_paragraph_rtl(paragraph: Paragraph) -> None:
    """``<w:bidi/>`` — paragraph-level right-to-left layout.

    This is what reverses the order of runs and puts the sentence-final full
    stop on the left. Alignment alone does not do it: a right-aligned
    left-to-right paragraph still reads in the wrong direction.
    """
    properties = paragraph._p.get_or_add_pPr()
    if properties.find(qn("w:bidi")) is None:
        properties.append(properties.makeelement(qn("w:bidi"), {}))


def _mark_run_rtl(run) -> None:
    """``<w:rtl/>`` — marks the run's characters as right-to-left.

    Needed in addition to ``w:bidi``: without it Word applies the *Latin* font
    to Arabic characters and picks its own substitute when that font has no
    Arabic glyphs.
    """
    properties = run._r.get_or_add_rPr()
    if properties.find(qn("w:rtl")) is None:
        properties.append(properties.makeelement(qn("w:rtl"), {}))


def _set_default_font(document) -> None:
    """Set both the Latin and the complex-script font on the Normal style.

    Word tracks them separately (``w:cs`` is the complex-script face), and both
    are set unconditionally rather than by the document's language: a single
    submission routinely carries Arabic answers and English standard names, so
    branching on one language leaves the other script to whatever the reader's
    machine happens to substitute — which is how a document renders differently
    on the evaluator's screen than on the bidder's.
    """
    style = document.styles["Normal"]
    style.font.size = Pt(11)

    # Created explicitly. Assigning `style.font.name` would create it as a side
    # effect, but it also writes the Latin face into the same element, which is
    # the branch this function exists to avoid.
    element = style.element.get_or_add_rPr().get_or_add_rFonts()
    element.set(qn("w:ascii"), _LATIN_FONT)
    element.set(qn("w:hAnsi"), _LATIN_FONT)
    element.set(qn("w:cs"), _ARABIC_FONT)


__all__ = ["MEDIA_TYPE", "build_response_docx", "is_rtl_text"]
